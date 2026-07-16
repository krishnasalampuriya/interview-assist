from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


MAX_UPLOAD_BYTES = 8 * 1024 * 1024
TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".json", ".rtf"}


class DocumentTextExtractionError(ValueError):
    pass


def extract_text_from_document(filename: str, content: bytes, content_type: str | None = None) -> str:
    if not content:
        raise DocumentTextExtractionError("Uploaded file is empty.")

    if len(content) > MAX_UPLOAD_BYTES:
        raise DocumentTextExtractionError("File is too large. Please upload a file up to 8 MB.")

    extension = Path(filename).suffix.lower()

    if extension == ".pdf" or content_type == "application/pdf":
        return _extract_pdf_text(content)

    if extension == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx_text(content)

    if extension in TEXT_EXTENSIONS or (content_type and content_type.startswith("text/")):
        return _extract_plain_text(content)

    if extension == ".doc":
        raise DocumentTextExtractionError("Old .doc files are not supported yet. Please upload PDF, DOCX, or TXT.")

    raise DocumentTextExtractionError("Unsupported file type. Please upload PDF, DOCX, TXT, MD, or CSV.")


def _extract_pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        page_text = [page.extract_text() or "" for page in reader.pages]
    except Exception as error:
        raise DocumentTextExtractionError("Could not read text from this PDF.") from error

    return _normalize_text("\n\n".join(page_text))


def _extract_docx_text(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as error:
        raise DocumentTextExtractionError("Could not read text from this DOCX file.") from error

    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return _normalize_text("\n".join(paragraphs))


def _extract_plain_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return _normalize_text(content.decode(encoding))
        except UnicodeDecodeError:
            continue

    raise DocumentTextExtractionError("Could not decode this text file.")


def _normalize_text(text: str) -> str:
    normalized = "\n".join(line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"))
    normalized = "\n".join(line for line in normalized.split("\n") if line)

    if not normalized:
        raise DocumentTextExtractionError("No readable text was found in this file.")

    return normalized

